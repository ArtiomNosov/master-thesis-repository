# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

root = Path('.')
cands = sorted(root.glob('otchet_po_praktike*.docx'))
if not cands:
    raise SystemExit('report file not found')
path = cands[0]

backup = path.with_name(path.stem + f"_backup_before_format_{datetime.now().strftime('%Y%m%d_%H%M%S')}" + path.suffix)
shutil.copy2(path, backup)


def set_run_font(run, size=12, bold=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, 12)


doc = Document(str(path))

# Page setup
for sec in doc.sections:
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1)
    sec.different_first_page_header_footer = True

# Base style
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
normal.font.size = Pt(12)
normal.font.color.rgb = RGBColor(0, 0, 0)
npf = normal.paragraph_format
npf.line_spacing = 1.5
npf.first_line_indent = Cm(1.25)
npf.space_before = Pt(0)
npf.space_after = Pt(0)
npf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Heading styles
for name, size in [('Heading 1', 14), ('Heading 2', 13), ('Heading 3', 12)]:
    st = doc.styles[name]
    st.font.name = 'Times New Roman'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    pf = st.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

h1_patterns = [
    r'^Содержание$', r'^Реферат$', r'^Введение$', r'^Заключение$',
    r'^Список использованных источников$', r'^Приложение\s+[А-ЯA-Z]'
]

# Clean paragraphs & apply formatting
prev_empty = False
for p in doc.paragraphs:
    text = p.text or ''
    text = re.sub(r' {2,}', ' ', text)
    p.text = text

    stripped = text.strip()
    if not stripped:
        if prev_empty:
            continue
        prev_empty = True
    else:
        prev_empty = False

    # infer heading levels when plain paragraph used
    if any(re.match(ptrn, stripped) for ptrn in h1_patterns) or re.match(r'^[1-9]\d*\.\s', stripped):
        if re.match(r'^\d+\.\d+\.\s', stripped):
            p.style = doc.styles['Heading 3']
        elif re.match(r'^\d+\.\d+\s', stripped):
            p.style = doc.styles['Heading 2']
        else:
            p.style = doc.styles['Heading 1']

    # caption handling
    if re.match(r'^(Таблица\s+\d+|Таблица\s+[0-9]+\.[0-9]+)', stripped):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
    elif re.match(r'^(Рис\.|Рисунок)\s*\d+', stripped):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
    elif p.style.name.startswith('Heading'):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)

    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    for run in p.runs:
        size = 12
        bold = run.bold
        if p.style.name == 'Heading 1':
            size, bold = 14, True
        elif p.style.name == 'Heading 2':
            size, bold = 13, True
        elif p.style.name == 'Heading 3':
            size, bold = 12, True
        set_run_font(run, size=size, bold=bold)

# Tables formatting
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=12)

# Footer page numbers (center), first page without number
for sec in doc.sections:
    footer = sec.footer
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    add_page_number(fp)

# Save
doc.save(str(path))
print(f'formatted: {path}')
print(f'backup: {backup}')
