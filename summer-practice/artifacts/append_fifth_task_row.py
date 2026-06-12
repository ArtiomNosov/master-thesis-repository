# -*- coding: utf-8 -*-
"""Добавляет 5-й пункт в таблицу заданий, не трогая остальные ячейки документа."""
from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

from docx import Document

DOC = Path(__file__).resolve().parent / "zadanie_letnyaya-praktika.docx"

NEW_TASK = (
    "Сводка результатов практики: сравнение метрик ранжирования с базовой линией, итоги "
    "интеграционных прогонов в контуре ATS; ограничения модели и приоритетные шаги к защите ВКР "
    "(данные, устойчивость/справедливость ранжирования, доработка пайплайна в experiments/)."
)
NEW_FORM = (
    "Сводные таблицы метрик, протоколы прогонов, перечень рисков и рекомендаций"
)
NEW_DEADLINE = "20.05.2026"


def set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    p0.add_run(text)


def main() -> int:
    if not DOC.is_file():
        print(f"Не найден: {DOC}", file=sys.stderr)
        return 1
    doc = Document(str(DOC))
    t3 = doc.tables[3]
    if len(t3.rows) > 5:
        print("Пятый пункт уже есть; строк в таблице:", len(t3.rows))
        return 0

    new_tr = deepcopy(t3.rows[-1]._tr)
    t3._tbl.append(new_tr)
    nr = t3.rows[-1]
    set_cell_text(nr.cells[1], NEW_TASK)
    set_cell_text(nr.cells[2], NEW_FORM)
    set_cell_text(nr.cells[3], NEW_DEADLINE)
    try:
        doc.save(str(DOC))
    except PermissionError:
        alt = DOC.with_name(DOC.stem + "_with_fifth_task.docx")
        doc.save(str(alt))
        print(
            "Не удалось записать (закройте файл в Word). Сохранено как:\n  ",
            alt,
            "\nПереименуйте в zadanie_letnyaya-praktika.docx при необходимости.",
        )
        return 0
    print("Добавлена строка задания. Всего строк в таблице (с заголовком):", len(t3.rows))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
