# -*- coding: utf-8 -*-
"""Fill MI ФИ Каф. №22 summer practice assignment docx from repo context."""
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "zadanie_letnyaya-praktika.docx"
OUT = ROOT / "zadanie_letnyaya-praktika.docx"

STUDENT_GROUP = "М24-534"
STUDENT_NAME = "Носов Артём Иванович"
SUPERVISOR = "Киреев Василий Сергеевич"
CONSULTANT = "Никифоров"

THEME = (
    "Продолжение работ по магистерской теме: разработка и летняя отладка "
    "интеллектуального модуля ранжирования кандидатов (семантический анализ "
    "резюме и вакансий, интеграция с контуром ATS) на базе НИЯУ МИФИ"
)

ROWS_TASKS = [
    (
        "Актуализировать постановку и обзор по теме ВКР с учётом результатов весны; "
        "подготовить разделы отчёта по практике.",
        "Тезисы, черновик отчёта",
        "15.07.2026",
    ),
    (
        "Доработать би-энкодер и пайплайн в experiments/: прогоны, сравнение метрик "
        "с весенней базовой линией, конфигурации и логи запусков.",
        "Код, отчёты, конфигурации",
        "15.08.2026",
    ),
    (
        f"Проверка ранжирования в тестовом контуре ATS; согласование с консультантом кафедры ({CONSULTANT}); протоколы проверок.",
        "Протоколы, метрики, материалы интеграции",
        "25.08.2026",
    ),
    (
        "Отчёт по практике; выводы для черновиков пояснительной записки и презентации.",
        "Отчёт по практике, фрагменты ПЗ и слайдов",
        "31.08.2026",
    ),
]

LITERATURE = [
    "Devlin J., Chang M. W., Lee K., Toutanova K. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding // Proceedings of NAACL-HLT 2019. — 2019. — P. 4171-4186.",
    "Reimers N., Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks // Proceedings of EMNLP-IJCNLP 2019. — 2019. — P. 3982-3992.",
    "Liu T. Y. Learning to Rank for Information Retrieval // Foundations and Trends in Information Retrieval. — 2009. — Vol. 3. — No. 3. — P. 225-331.",
    "Jurafsky D., Martin J. H. Speech and Language Processing // 3rd ed. draft. — Stanford University, 2024.",
    "Manning C. D., Raghavan P., Schütze H. Introduction to Information Retrieval. — Cambridge University Press, 2008.",
]


def set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    p0.add_run(text)


def delete_table_last_rows(table, target_row_count: int) -> None:
    tbl = table._tbl
    excess = len(table.rows) - target_row_count
    for _ in range(excess):
        tbl.remove(table.rows[-1]._tr)


def main() -> None:
    doc = Document(str(SRC))

    t1 = doc.tables[1]
    set_cell_text(t1.rows[0].cells[1], STUDENT_GROUP)
    set_cell_text(t1.rows[0].cells[3], STUDENT_NAME)

    t2 = doc.tables[2]
    set_cell_text(t2.rows[0].cells[0], THEME)

    t3 = doc.tables[3]
    want_total = 1 + len(ROWS_TASKS)
    for i, (content, form, deadline) in enumerate(ROWS_TASKS):
        r = t3.rows[i + 1]
        set_cell_text(r.cells[1], content)
        set_cell_text(r.cells[2], form)
        if len(r.cells) > 3 and deadline:
            set_cell_text(r.cells[3], deadline)
    delete_table_last_rows(t3, want_total)

    t4 = doc.tables[4]
    for i, ref in enumerate(LITERATURE):
        if i >= len(t4.rows):
            break
        set_cell_text(t4.rows[i].cells[1], ref)

    t5 = doc.tables[5]
    set_cell_text(t5.rows[2].cells[1], "19")
    set_cell_text(t5.rows[2].cells[3], "мая")
    set_cell_text(t5.rows[2].cells[4], "2026 г.")
    set_cell_text(t5.rows[0].cells[7], SUPERVISOR)
    set_cell_text(t5.rows[2].cells[8], STUDENT_NAME)

    doc.save(str(OUT))


if __name__ == "__main__":
    main()
