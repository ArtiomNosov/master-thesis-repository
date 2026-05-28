#!/usr/bin/env python3
"""Heuristic DOCX review report for obvious formatting issues (read-only)."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", required=True, type=Path)
    args = parser.parse_args()

    document = Document(str(args.docx))
    issues: list[str] = []

    for index, paragraph in enumerate(document.paragraphs, 1):
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name == "Normal" and text.startswith("**") and text.endswith("**"):
            issues.append(f"П.{index}: возможный нераспознанный заголовок в Normal: {text[:80]}")
        if re.search(r"\$[0-9]", text):
            issues.append(f"П.{index}: осталась оценка в USD: {text[:100]}")
        if text.startswith("@startuml") or (paragraph.style.name == "VKR Code Block" and "@" in text):
            issues.append(f"П.{index}: PlantUML/код в теле работы без подписи «Рисунок/Листинг»")
        if re.match(r"^\[\d+\]", text) and paragraph.style.name not in {"VKR Bibliography"}:
            issues.append(f"П.{index}: элемент списка литературы не в стиле VKR Bibliography")
        if "Estimate cost" in text or "Pair-level classification" in text:
            issues.append(f"П.{index}: англоязычный заголовок/фрагмент в таблице или тексте")

    table_count = 0
    for table in document.tables:
        if "<m:oMath" in table._tbl.xml:
            widths = [column.width for column in table.columns]
            if len(widths) >= 2 and widths[0] < widths[-1]:
                issues.append(f"Формула-таблица {table_count + 1}: узкая колонка уравнения ({widths})")
            continue
        table_count += 1

    if not any("<m:oMath" in table._tbl.xml for table in document.tables):
        issues.append("В документе не найдены нативные формулы Word (m:oMath)")

    print("=== Обзор очевидных проблем форматирования ===")
    if not issues:
        print("Явных проблем не найдено эвристикой скрипта.")
        return
    for item in issues:
        print(f"- {item}")


if __name__ == "__main__":
    main()
