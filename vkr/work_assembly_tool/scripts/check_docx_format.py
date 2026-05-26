#!/usr/bin/env python3
"""Check generated DOCX formatting against the VКР JSON config."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH


ALIGNMENT_NAMES = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}


def approx(actual: float | None, expected: float, tolerance: float = 0.05) -> bool:
    return actual is not None and abs(actual - expected) <= tolerance


def load_config(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def fail(errors: list[str], message: str) -> None:
    errors.append(message)


def warn(warnings: list[str], message: str) -> None:
    warnings.append(message)


def check_section(document: Document, config: dict[str, Any], errors: list[str]) -> None:
    section = document.sections[0]
    page = config["document"]["page"]
    margins = page["margins_mm"]
    checks = {
        "page_width_mm": (section.page_width.mm, page["width_mm"]),
        "page_height_mm": (section.page_height.mm, page["height_mm"]),
        "top_margin_mm": (section.top_margin.mm, margins["top"]),
        "bottom_margin_mm": (section.bottom_margin.mm, margins["bottom"]),
        "left_margin_mm": (section.left_margin.mm, margins["left"]),
        "right_margin_mm": (section.right_margin.mm, margins["right"]),
    }
    for name, (actual, expected) in checks.items():
        if not approx(actual, expected, tolerance=0.2):
            fail(errors, f"{name}: expected {expected}, got {actual}")


def font_size_pt(style: Any) -> float | None:
    return style.font.size.pt if style.font.size is not None else None


def check_style(document: Document, style_name: str, expected: dict[str, Any], errors: list[str]) -> None:
    style = document.styles[style_name]
    if "font_size_pt" in expected and not approx(font_size_pt(style), expected["font_size_pt"]):
        fail(errors, f"{style_name}: font size expected {expected['font_size_pt']}, got {font_size_pt(style)}")
    if "bold" in expected and style.font.bold != expected["bold"]:
        fail(errors, f"{style_name}: bold expected {expected['bold']}, got {style.font.bold}")
    paragraph_format = style.paragraph_format
    if "line_spacing" in expected and not approx(float(paragraph_format.line_spacing), expected["line_spacing"]):
        fail(errors, f"{style_name}: line spacing expected {expected['line_spacing']}, got {paragraph_format.line_spacing}")
    if "first_line_indent_cm" in expected:
        actual = paragraph_format.first_line_indent.cm if paragraph_format.first_line_indent is not None else 0
        if not approx(actual, expected["first_line_indent_cm"]):
            fail(errors, f"{style_name}: first line indent expected {expected['first_line_indent_cm']}, got {actual}")
    if "space_before_pt" in expected:
        actual = paragraph_format.space_before.pt if paragraph_format.space_before is not None else 0
        if not approx(actual, expected["space_before_pt"]):
            fail(errors, f"{style_name}: space before expected {expected['space_before_pt']}, got {actual}")
    if "space_after_pt" in expected:
        actual = paragraph_format.space_after.pt if paragraph_format.space_after is not None else 0
        if not approx(actual, expected["space_after_pt"]):
            fail(errors, f"{style_name}: space after expected {expected['space_after_pt']}, got {actual}")


def check_styles(document: Document, config: dict[str, Any], errors: list[str]) -> None:
    default = config["document"]["default_text"]
    check_style(document, "Normal", {"font_size_pt": default["font_size_pt"], **default}, errors)
    for level, values in config["headings"]["levels"].items():
        check_style(document, f"Heading {level}", values, errors)
    check_style(document, "VKR Table Caption", config["captions"]["table"], errors)
    check_style(document, "VKR Figure Caption", config["captions"]["figure"], errors)
    check_style(document, "VKR Code Block", config["code_blocks"], errors)


def color_to_hex(color: Any) -> str | None:
    return str(color.rgb) if color is not None and color.rgb is not None else None


def check_text_color(document: Document, config: dict[str, Any], errors: list[str]) -> None:
    expected = config["document"]["default_text"]["font_color"].upper()
    for style in document.styles:
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
            color = color_to_hex(style.font.color)
            if color is not None and color.upper() != expected:
                fail(errors, f"Style {style.name}: text color expected {expected}, got {color}")
    for paragraph_index, paragraph in enumerate(document.paragraphs, 1):
        for run_index, run in enumerate(paragraph.runs, 1):
            if not run.text.strip():
                continue
            color = color_to_hex(run.font.color)
            if color is not None and color.upper() != expected:
                fail(errors, f"Paragraph {paragraph_index} run {run_index}: text color expected {expected}, got {color}")
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            for cell_index, cell in enumerate(row.cells, 1):
                for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
                    for run_index, run in enumerate(paragraph.runs, 1):
                        if not run.text.strip():
                            continue
                        color = color_to_hex(run.font.color)
                        if color is not None and color.upper() != expected:
                            fail(
                                errors,
                                f"Table {table_index} row {row_index} cell {cell_index} paragraph {paragraph_index} run {run_index}: text color expected {expected}, got {color}",
                            )


def paragraph_effective_alignment(paragraph: Any) -> str | None:
    alignment = paragraph.alignment
    if alignment is None:
        alignment = paragraph.style.paragraph_format.alignment
    return ALIGNMENT_NAMES.get(alignment)


def check_paragraphs(document: Document, config: dict[str, Any], errors: list[str]) -> None:
    empty_allowed = config["quality_checks"]["empty_paragraphs_allowed"]
    double_spaces_allowed = config["quality_checks"]["double_spaces_allowed"]
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = paragraph.text
        style_name = paragraph.style.name
        if not empty_allowed and not text.strip():
            fail(errors, f"Paragraph {index}: empty paragraph is not allowed")
        if not double_spaces_allowed and style_name != "VKR Code Block" and re.search(r" {2,}", text):
            fail(errors, f"Paragraph {index}: double spaces found")
        if re.match(r"^Таблица\s+[\w.]+\s+[-–]\s+", text):
            if style_name != "VKR Table Caption":
                fail(errors, f"Paragraph {index}: table caption has style {style_name}")
            if paragraph_effective_alignment(paragraph) != config["captions"]["table"]["alignment"]:
                fail(errors, f"Paragraph {index}: table caption alignment mismatch")
        if re.match(r"^Рисунок\s+[\w.]+\s+[-–]\s+", text):
            if style_name != "VKR Figure Caption":
                fail(errors, f"Paragraph {index}: figure caption has style {style_name}")
            if paragraph_effective_alignment(paragraph) != config["captions"]["figure"]["alignment"]:
                fail(errors, f"Paragraph {index}: figure caption alignment mismatch")


def check_tables(document: Document, config: dict[str, Any], errors: list[str]) -> None:
    expected_spacing = config["tables"]["cell_line_spacing"]
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            for cell_index, cell in enumerate(row.cells, 1):
                for paragraph in cell.paragraphs:
                    if paragraph.paragraph_format.line_spacing is not None:
                        actual = float(paragraph.paragraph_format.line_spacing)
                        if not approx(actual, expected_spacing):
                            fail(
                                errors,
                                f"Table {table_index} row {row_index} cell {cell_index}: line spacing expected {expected_spacing}, got {actual}",
                            )


def iter_body_blocks(document: Document) -> list[tuple[str, Any]]:
    paragraph_by_id = {id(paragraph._p): paragraph for paragraph in document.paragraphs}
    table_by_id = {id(table._tbl): table for table in document.tables}
    blocks: list[tuple[str, Any]] = []
    for child in document.element.body.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if tag == "p" and id(child) in paragraph_by_id:
            blocks.append(("paragraph", paragraph_by_id[id(child)]))
        elif tag == "tbl" and id(child) in table_by_id:
            blocks.append(("table", table_by_id[id(child)]))
    return blocks


def is_table_caption_text(text: str, config: dict[str, Any]) -> bool:
    prefix = config["captions"]["table"]["prefix"]
    stripped = text.strip()
    return bool(re.match(rf"^{prefix}\s+[\w.]+\s+[-–]\s+\S+", stripped, re.IGNORECASE))


def is_figure_caption_text(text: str, config: dict[str, Any]) -> bool:
    prefix = config["captions"]["figure"]["prefix"]
    stripped = text.strip()
    return bool(re.match(rf"^{prefix}\s+[\w.]+\s+[-–]\s+\S+", stripped, re.IGNORECASE))


def previous_nonempty_paragraph(blocks: list[tuple[str, Any]], start_index: int) -> Any | None:
    for kind, block in reversed(blocks[:start_index]):
        if kind == "paragraph" and block.text.strip():
            return block
    return None


def next_nonempty_paragraph(blocks: list[tuple[str, Any]], start_index: int) -> Any | None:
    for kind, block in blocks[start_index + 1:]:
        if kind == "paragraph" and block.text.strip():
            return block
    return None


def paragraph_has_drawing(paragraph: Any) -> bool:
    return "<w:drawing" in paragraph._p.xml or "<w:pict" in paragraph._p.xml


def check_required_captions(document: Document, config: dict[str, Any], errors: list[str]) -> None:
    blocks = iter_body_blocks(document)
    table_number = 0
    figure_number = 0
    for index, (kind, block) in enumerate(blocks):
        if kind == "table":
            table_number += 1
            caption = previous_nonempty_paragraph(blocks, index)
            if caption is None or caption.style.name != "VKR Table Caption" or not is_table_caption_text(caption.text, config):
                fail(errors, f"Table {table_number}: missing required caption above the table")
        elif kind == "paragraph" and paragraph_has_drawing(block):
            figure_number += 1
            caption = next_nonempty_paragraph(blocks, index)
            if caption is None or caption.style.name != "VKR Figure Caption" or not is_figure_caption_text(caption.text, config):
                fail(errors, f"Figure {figure_number}: missing required caption below the figure")


def check_page_number(document: Document, config: dict[str, Any], errors: list[str]) -> None:
    if not config["document"]["page_numbers"]["enabled"]:
        return
    footer_xml = document.sections[0].footer._element.xml
    if "PAGE" not in footer_xml:
        fail(errors, "Footer PAGE field is missing")


def count_references(document: Document) -> int:
    text = "\n".join(p.text for p in document.paragraphs)
    return len(set(re.findall(r"\[(\d+)\]", text)))


def check_reference_counts(document: Document, config: dict[str, Any], warnings: list[str], strict: bool, errors: list[str]) -> None:
    reference_rules = config["quality_checks"]["references"]
    total = count_references(document)
    message = f"Reference citation count by [n] markers is {total}; expected at least {reference_rules['min_total']}"
    if total < reference_rules["min_total"]:
        if strict:
            fail(errors, message)
        else:
            warn(warnings, message)


def run_checks(docx_path: Path, config: dict[str, Any], strict_content_checks: bool) -> tuple[list[str], list[str]]:
    document = Document(str(docx_path))
    errors: list[str] = []
    warnings: list[str] = []
    check_section(document, config, errors)
    check_styles(document, config, errors)
    check_text_color(document, config, errors)
    check_paragraphs(document, config, errors)
    check_tables(document, config, errors)
    check_required_captions(document, config, errors)
    check_page_number(document, config, errors)
    check_reference_counts(document, config, warnings, strict_content_checks, errors)
    return errors, warnings


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", required=True, type=Path, help="Generated DOCX file")
    parser.add_argument("--config", required=True, type=Path, help="VКР formatting JSON config")
    parser.add_argument("--strict-content-checks", action="store_true", help="Fail on content count checks")
    args = parser.parse_args()

    errors, warnings = run_checks(args.docx, load_config(args.config), args.strict_content_checks)
    for message in warnings:
        print(f"WARN: {message}")
    if errors:
        for message in errors:
            print(f"ERROR: {message}", file=sys.stderr)
        raise SystemExit(1)
    print("DOCX formatting checks passed")


if __name__ == "__main__":
    main()
