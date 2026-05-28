#!/usr/bin/env python3
"""Convert a Markdown thesis draft to a DOCX formatted by VКР rules."""

from __future__ import annotations

import argparse
import json
import re
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docxcompose.composer import Composer
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.token import Token

from formula_omml import formula_text_to_omml, looks_like_formula, set_table_no_borders
from plantuml_render import extract_plantuml_title, render_plantuml_png_stream


EXPECTED_CONFIG_KEYS: dict[str, Any] = {
    "document": {
        "page": {"size", "orientation", "width_mm", "height_mm", "margins_mm"},
        "default_text": {
            "font_family",
            "font_size_pt",
            "font_color",
            "alignment",
            "line_spacing",
            "first_line_indent_cm",
            "space_before_pt",
            "space_after_pt",
        },
        "page_numbers": {"enabled", "footer_alignment"},
    },
    "headings": {"max_level_in_toc", "level_1_starts_new_page", "levels"},
    "captions": {"table", "figure", "listing"},
    "tables": {"cell_line_spacing", "cell_space_before_pt", "cell_space_after_pt", "cell_margin_dxa", "repeat_header_rows"},
    "formulas": {"alignment", "number_alignment", "number_enclosure", "number_column_cm"},
    "code_blocks": {"line_spacing", "space_before_pt", "space_after_pt", "first_line_indent_cm"},
    "bibliography": {"line_spacing", "left_indent_cm", "hanging_indent_cm", "space_after_pt"},
    "markdown": {"encoding", "promote_full_strong_paragraphs_to_headings", "supported_block_types"},
    "quality_checks": {"empty_paragraphs_allowed", "double_spaces_allowed", "references"},
    "assembly": {
        "title_pages",
        "abstract_md",
        "include_toc",
        "toc_heading",
        "toc_levels",
        "plantuml_server_url",
    },
}

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _validate_known_keys(data: dict[str, Any], expected: dict[str, Any] | set[str], path: str = "") -> None:
    if isinstance(expected, set):
        unknown = set(data) - expected
        if unknown:
            raise ValueError(f"Unknown config keys at {path or '<root>'}: {sorted(unknown)}")
        return

    unknown = set(data) - set(expected)
    if unknown:
        raise ValueError(f"Unknown config keys at {path or '<root>'}: {sorted(unknown)}")
    for key, nested in expected.items():
        if key not in data:
            raise ValueError(f"Missing config key: {path + '.' if path else ''}{key}")
        if isinstance(nested, (dict, set)):
            if not isinstance(data[key], dict):
                raise ValueError(f"Config key must be an object: {path + '.' if path else ''}{key}")
            _validate_known_keys(data[key], nested, path + "." + key if path else key)


def load_config(path: Path) -> dict[str, Any]:
    config = json.loads(path.read_text(encoding="utf-8"))
    _validate_known_keys(config, EXPECTED_CONFIG_KEYS)
    return config


def set_russian_font(style_or_run: Any, family: str) -> None:
    font = style_or_run.font
    font.name = family
    rpr = style_or_run.element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        style_or_run.element.insert(0, rpr)
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), family)


def apply_paragraph_format(paragraph_format: Any, values: dict[str, Any]) -> None:
    if "alignment" in values:
        paragraph_format.alignment = ALIGNMENTS[values["alignment"]]
    if "line_spacing" in values:
        paragraph_format.line_spacing = values["line_spacing"]
    if "first_line_indent_cm" in values:
        paragraph_format.first_line_indent = Cm(values["first_line_indent_cm"])
    if "space_before_pt" in values:
        paragraph_format.space_before = Pt(values["space_before_pt"])
    if "space_after_pt" in values:
        paragraph_format.space_after = Pt(values["space_after_pt"])
    if "keep_with_next" in values:
        paragraph_format.keep_with_next = bool(values["keep_with_next"])


def force_text_styles_black(document: Document, color: str) -> None:
    rgb = RGBColor.from_string(color)
    for style in document.styles:
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
            clear_style_theme_color(style)
            style.font.color.rgb = rgb


def set_run_text_black(run: Any, config: dict[str, Any]) -> None:
    run.font.color.rgb = RGBColor.from_string(config["document"]["default_text"]["font_color"])


def apply_run_font(run: Any, config: dict[str, Any]) -> None:
    family = config["document"]["default_text"]["font_family"]
    set_russian_font(run, family)
    set_run_text_black(run, config)


def apply_paragraph_run_fonts(paragraph: Any, config: dict[str, Any]) -> None:
    for run in paragraph.runs:
        apply_run_font(run, config)


def clear_style_theme_color(style: Any) -> None:
    rpr = style.element.rPr
    if rpr is None:
        return
    color = rpr.find(qn("w:color"))
    if color is not None:
        if color.get(qn("w:themeColor")) is not None:
            del color.attrib[qn("w:themeColor")]
        if color.get(qn("w:themeTint")) is not None:
            del color.attrib[qn("w:themeTint")]
        if color.get(qn("w:themeShade")) is not None:
            del color.attrib[qn("w:themeShade")]


def configure_document(document: Document, config: dict[str, Any]) -> None:
    page = config["document"]["page"]
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT if page["orientation"] == "portrait" else WD_ORIENT.LANDSCAPE
    section.page_width = Mm(page["width_mm"])
    section.page_height = Mm(page["height_mm"])
    margins = page["margins_mm"]
    section.top_margin = Mm(margins["top"])
    section.bottom_margin = Mm(margins["bottom"])
    section.left_margin = Mm(margins["left"])
    section.right_margin = Mm(margins["right"])

    default = config["document"]["default_text"]
    normal = document.styles["Normal"]
    set_russian_font(normal, default["font_family"])
    normal.font.size = Pt(default["font_size_pt"])
    normal.font.color.rgb = RGBColor.from_string(default["font_color"])
    apply_paragraph_format(normal.paragraph_format, default)

    for level_text, heading_config in config["headings"]["levels"].items():
        level = int(level_text)
        style = document.styles[f"Heading {min(level, 9)}"]
        set_russian_font(style, default["font_family"])
        style.font.size = Pt(heading_config["font_size_pt"])
        style.font.bold = bool(heading_config["bold"])
        apply_paragraph_format(style.paragraph_format, heading_config)
        if level == 1:
            style.paragraph_format.page_break_before = bool(config["headings"]["level_1_starts_new_page"])

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        set_russian_font(style, default["font_family"])
        style.font.size = Pt(default["font_size_pt"])
        style.paragraph_format.line_spacing = default["line_spacing"]
        style.paragraph_format.space_before = Pt(default["space_before_pt"])
        style.paragraph_format.space_after = Pt(default["space_after_pt"])

    create_or_update_caption_style(document, "VKR Table Caption", config["captions"]["table"], default)
    create_or_update_caption_style(document, "VKR Figure Caption", config["captions"]["figure"], default)
    create_or_update_caption_style(document, "VKR Listing Caption", config["captions"]["listing"], default)
    create_or_update_code_style(document, config["code_blocks"], default)
    create_or_update_bibliography_style(document, config["bibliography"], default)
    force_text_styles_black(document, default["font_color"])

    if config["document"]["page_numbers"]["enabled"]:
        add_page_number(section.footer, config["document"]["page_numbers"]["footer_alignment"])


def create_or_update_caption_style(
    document: Document, style_name: str, caption_config: dict[str, Any], default: dict[str, Any]
) -> None:
    styles = document.styles
    style = styles[style_name] if style_name in styles else styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    set_russian_font(style, default["font_family"])
    style.font.size = Pt(caption_config["font_size_pt"])
    style.font.color.rgb = RGBColor.from_string(default["font_color"])
    apply_paragraph_format(style.paragraph_format, caption_config)


def create_or_update_bibliography_style(
    document: Document, bibliography_config: dict[str, Any], default: dict[str, Any]
) -> None:
    styles = document.styles
    style_name = "VKR Bibliography"
    style = styles[style_name] if style_name in styles else styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    set_russian_font(style, default["font_family"])
    style.font.size = Pt(default["font_size_pt"])
    style.font.color.rgb = RGBColor.from_string(default["font_color"])
    style.paragraph_format.line_spacing = bibliography_config["line_spacing"]
    style.paragraph_format.left_indent = Cm(bibliography_config["left_indent_cm"])
    style.paragraph_format.first_line_indent = Cm(bibliography_config["hanging_indent_cm"])
    style.paragraph_format.space_after = Pt(bibliography_config["space_after_pt"])
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def length_to_twips(value: Any) -> int:
    if hasattr(value, "twips"):
        return int(value.twips)
    if hasattr(value, "inches"):
        return int(value.inches * 1440)
    return int(value)


def set_table_preferred_width(table: Any, width: Any) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(length_to_twips(width)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def wrap_omath_para(omath_element: Any) -> Any:
    o_math_para = OxmlElement("m:oMathPara")
    o_math_para_pr = OxmlElement("m:oMathParaPr")
    jc = OxmlElement("m:jc")
    jc.set(qn("m:val"), "center")
    o_math_para_pr.append(jc)
    o_math_para.append(o_math_para_pr)
    o_math_para.append(omath_element)
    return o_math_para


def create_or_update_code_style(document: Document, code_config: dict[str, Any], default: dict[str, Any]) -> None:
    styles = document.styles
    style_name = "VKR Code Block"
    style = styles[style_name] if style_name in styles else styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    set_russian_font(style, default["font_family"])
    style.font.size = Pt(default["font_size_pt"])
    style.font.color.rgb = RGBColor.from_string(default["font_color"])
    apply_paragraph_format(style.paragraph_format, code_config)


def chapter_from_heading_text(text: str) -> int | None:
    match = re.match(r"^(\d+)\.", text.strip())
    return int(match.group(1)) if match else None


def next_figure_label(chapter: int | None, chapter_figures: dict[int, int]) -> str:
    if chapter is None:
        chapter = 0
        chapter_figures.setdefault(0, 0)
    chapter_figures[chapter] = chapter_figures.get(chapter, 0) + 1
    if chapter == 0:
        return str(chapter_figures[chapter])
    return f"{chapter}.{chapter_figures[chapter]}"


def add_figure_caption(document: Document, label: str, title: str, config: dict[str, Any]) -> None:
    figure_cfg = config["captions"]["figure"]
    paragraph = document.add_paragraph(style="VKR Figure Caption")
    paragraph.text = f"{figure_cfg['prefix']} {label} {figure_cfg['separator']} {title}"
    for run in paragraph.runs:
        apply_run_font(run, config)


def add_plantuml_figure(
    document: Document,
    content: str,
    config: dict[str, Any],
    chapter: int | None,
    chapter_figures: dict[int, int],
) -> None:
    assembly = config.get("assembly", {})
    server_url = assembly.get("plantuml_server_url", "http://www.plantuml.com/plantuml/img/")
    png_stream = render_plantuml_png_stream(content, server_url)
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    max_width = Cm(config["captions"]["figure"].get("max_width_cm", 15.0))
    run.add_picture(png_stream, width=max_width)

    title = extract_plantuml_title(content) or "Схема"
    label = next_figure_label(chapter, chapter_figures)
    add_figure_caption(document, label, title, config)


def is_listing_caption_text(text: str, config: dict[str, Any]) -> bool:
    prefix = config["captions"]["listing"]["prefix"]
    return bool(re.match(rf"^{prefix}\s+[\w.]+\s+[-–]\s+\S+", text.strip(), re.IGNORECASE))


def add_toc_page(document: Document, config: dict[str, Any]) -> None:
    assembly = config.get("assembly", {})
    heading = assembly.get("toc_heading", "Содержание")
    levels = assembly.get("toc_levels", "1-3")

    heading_paragraph = document.add_heading(level=1)
    heading_paragraph.text = ""
    run = heading_paragraph.add_run(heading)
    apply_run_font(run, config)
    heading_paragraph.paragraph_format.first_line_indent = Cm(0)

    toc_paragraph = document.add_paragraph(style="Normal")
    toc_paragraph.paragraph_format.first_line_indent = Cm(0)
    toc_run = toc_paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f'TOC \\o "{levels}" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    toc_run._r.append(fld_begin)
    toc_run._r.append(instr_text)
    toc_run._r.append(fld_separate)
    toc_run._r.append(fld_end)
    document.add_page_break()


def ensure_page_numbers(document: Document, config: dict[str, Any]) -> None:
    if not config["document"]["page_numbers"]["enabled"]:
        return
    alignment = config["document"]["page_numbers"]["footer_alignment"]
    for section in document.sections:
        add_page_number(section.footer, alignment)


def compose_documents(parts: list[Path], output_path: Path, config: dict[str, Any] | None = None) -> None:
    if not parts:
        raise ValueError("No documents to compose")
    master = Document(str(parts[0]))
    composer = Composer(master)
    for part in parts[1:]:
        composer.append(Document(str(part)))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(output_path))
    if config is not None:
        final_document = Document(str(output_path))
        ensure_page_numbers(final_document, config)
        final_document.save(str(output_path))


def resolve_tool_path(config_path: Path, relative_path: str) -> Path:
    candidate = (config_path.parent / relative_path).resolve()
    if not candidate.exists():
        raise FileNotFoundError(f"Assembly file not found: {candidate}")
    return candidate


def assemble_vkr_document(
    body_path: Path,
    output_path: Path,
    config_path: Path,
    config: dict[str, Any],
    abstract_path: Path | None,
) -> None:
    assembly = config.get("assembly", {})
    parts: list[Path] = []
    for relative in assembly.get("title_pages", []):
        parts.append(resolve_tool_path(config_path, relative))
    if abstract_path is not None:
        parts.append(abstract_path)
    if assembly.get("include_toc", False):
        toc_doc = Document()
        configure_document(toc_doc, config)
        add_toc_page(toc_doc, config)
        fd, toc_temp_name = tempfile.mkstemp(suffix="_toc.docx")
        import os

        os.close(fd)
        toc_temp = Path(toc_temp_name)
        toc_doc.save(str(toc_temp))
        parts.append(toc_temp)
    parts.append(body_path)
    compose_documents(parts, output_path, config)


def add_page_number(footer: Any, alignment: str) -> None:
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = ALIGNMENTS[alignment]
    paragraph.text = ""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def normalize_text(text: str, config: dict[str, Any]) -> str:
    if not config["quality_checks"]["double_spaces_allowed"]:
        text = re.sub(r"[ \t]{2,}", " ", text)
    return text


def plain_text_from_inline(token: Token) -> str:
    if not token.children:
        return token.content
    parts: list[str] = []
    for child in token.children:
        if child.type in {"text", "code_inline"}:
            parts.append(child.content)
        elif child.type in {"softbreak", "hardbreak"}:
            parts.append("\n")
        elif child.type == "image":
            parts.append(child.content)
    return "".join(parts)


def full_strong_text(token: Token) -> str | None:
    children = token.children or []
    meaningful = [child for child in children if not (child.type == "text" and not child.content.strip())]
    if len(meaningful) == 3 and meaningful[0].type == "strong_open" and meaningful[2].type == "strong_close":
        if meaningful[1].type == "text":
            return meaningful[1].content.strip()
    return None


def promoted_heading_level(text: str) -> int | None:
    normalized = text.strip()
    if re.match(r"^(Введение|Заключение|Список\b|Содержание\b)", normalized, re.IGNORECASE):
        return 1
    numbered = re.match(r"^(\d+(?:\.\d+)*)\.?\s+\S", normalized)
    if numbered:
        return min(numbered.group(1).count(".") + 1, 3)
    if len(normalized) <= 140 and not normalized.endswith("."):
        return 2
    return None


def add_inline_runs(paragraph: Any, token: Token, config: dict[str, Any]) -> None:
    if not token.children:
        run = paragraph.add_run(normalize_text(token.content, config))
        set_run_text_black(run, config)
        return

    active = {"bold": False, "italic": False, "code": False}
    link_depth = 0
    for child in token.children:
        if child.type == "strong_open":
            active["bold"] = True
        elif child.type == "strong_close":
            active["bold"] = False
        elif child.type == "em_open":
            active["italic"] = True
        elif child.type == "em_close":
            active["italic"] = False
        elif child.type == "code_inline":
            run = paragraph.add_run(normalize_text(child.content, config))
            run.bold = active["bold"]
            run.italic = active["italic"]
            apply_run_font(run, config)
        elif child.type == "softbreak":
            run = paragraph.add_run(" ")
            set_run_text_black(run, config)
        elif child.type == "hardbreak":
            paragraph.add_run().add_break()
        elif child.type == "link_open":
            link_depth += 1
        elif child.type == "link_close":
            link_depth = max(0, link_depth - 1)
        elif child.type == "image":
            src = dict(child.attrs or {}).get("src")
            if src and Path(src).exists():
                paragraph._parent.add_picture(src)
            elif child.content:
                run = paragraph.add_run(normalize_text(child.content, config))
                set_run_text_black(run, config)
        elif child.type == "text":
            run = paragraph.add_run(normalize_text(child.content, config))
            run.bold = active["bold"]
            run.italic = active["italic"]
            set_run_text_black(run, config)
            if link_depth:
                run.underline = True


def add_paragraph_from_inline(
    document: Document,
    token: Token,
    config: dict[str, Any],
    list_style: str | None,
    ordered_number: int | None = None,
    in_bibliography: list[bool] | None = None,
) -> None:
    if in_bibliography is None:
        in_bibliography = [False]
    text = plain_text_from_inline(token).strip()
    if not text and not config["quality_checks"]["empty_paragraphs_allowed"]:
        return

    if in_bibliography[0] and re.match(r"^\[\d+\]", text):
        paragraph = document.add_paragraph(style="VKR Bibliography")
        bibliography = config["bibliography"]
        paragraph.paragraph_format.left_indent = Cm(bibliography["left_indent_cm"])
        paragraph.paragraph_format.first_line_indent = Cm(bibliography["hanging_indent_cm"])
        paragraph.paragraph_format.line_spacing = bibliography["line_spacing"]
        add_inline_runs(paragraph, token, config)
        apply_paragraph_run_fonts(paragraph, config)
        return

    if ordered_number is not None:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.left_indent = Cm(1.25)
        paragraph.paragraph_format.first_line_indent = Cm(-0.635)
        prefix = paragraph.add_run(f"{ordered_number}.\t")
        apply_run_font(prefix, config)
        add_inline_runs(paragraph, token, config)
        apply_paragraph_run_fonts(paragraph, config)
        return

    if list_style:
        paragraph = document.add_paragraph(style=list_style)
        add_inline_runs(paragraph, token, config)
        apply_paragraph_run_fonts(paragraph, config)
        return

    table_caption_prefix = config["captions"]["table"]["prefix"]
    figure_caption_prefix = config["captions"]["figure"]["prefix"]
    listing_caption_prefix = config["captions"]["listing"]["prefix"]
    if re.match(rf"^(Продолжение|Окончание)\s+{table_caption_prefix}\b", text, re.IGNORECASE):
        paragraph = document.add_paragraph(style="VKR Table Caption")
    elif re.match(rf"^{table_caption_prefix}\s+[\w.]+\s+[-–]\s+", text, re.IGNORECASE):
        paragraph = document.add_paragraph(style="VKR Table Caption")
    elif re.match(rf"^{figure_caption_prefix}\s+[\w.]+\s+[-–]\s+", text, re.IGNORECASE):
        paragraph = document.add_paragraph(style="VKR Figure Caption")
    elif re.match(rf"^{listing_caption_prefix}\s+[\w.]+\s+[-–]\s+", text, re.IGNORECASE):
        paragraph = document.add_paragraph(style="VKR Listing Caption")
    else:
        strong_text = full_strong_text(token)
        if strong_text and config["markdown"]["promote_full_strong_paragraphs_to_headings"]:
            level = promoted_heading_level(strong_text)
            if level:
                paragraph = document.add_heading(level=level)
                paragraph.text = ""
                run = paragraph.add_run(normalize_text(strong_text, config))
                apply_run_font(run, config)
                paragraph.paragraph_format.first_line_indent = Cm(0)
                if re.match(r"^Список литературы", strong_text, re.IGNORECASE):
                    in_bibliography[0] = True
                return
        paragraph = document.add_paragraph(style="Normal")

    add_inline_runs(paragraph, token, config)
    apply_paragraph_run_fonts(paragraph, config)


def add_heading(document: Document, token: Token, level: int, config: dict[str, Any]) -> None:
    paragraph = document.add_heading(level=min(level, 9))
    paragraph.text = ""
    add_inline_runs(paragraph, token, config)
    apply_paragraph_run_fonts(paragraph, config)
    paragraph.paragraph_format.first_line_indent = Cm(0)


def set_cell_margins(cell: Any, margin_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:w"), str(margin_dxa))
        element.set(qn("w:type"), "dxa")
        tc_mar.append(element)
    tc_pr.append(tc_mar)


def format_table(table: Any, config: dict[str, Any], default: dict[str, Any]) -> None:
    margin_dxa = config["tables"]["cell_margin_dxa"]
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, margin_dxa)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.left_indent = Cm(0)
                paragraph.paragraph_format.right_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = config["tables"]["cell_line_spacing"]
                paragraph.paragraph_format.space_before = Pt(config["tables"]["cell_space_before_pt"])
                paragraph.paragraph_format.space_after = Pt(config["tables"]["cell_space_after_pt"])
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    apply_run_font(run, config)
                    run.font.size = Pt(default["font_size_pt"])


def set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def is_table_caption_text(text: str, config: dict[str, Any]) -> bool:
    table_caption_prefix = config["captions"]["table"]["prefix"]
    stripped = text.strip()
    return bool(
        re.match(rf"^{table_caption_prefix}\s+[\w.]+\s+[-–]\s+\S+", stripped, re.IGNORECASE)
        or re.match(rf"^(Продолжение|Окончание)\s+{table_caption_prefix}\b", stripped, re.IGNORECASE)
    )


def document_ends_with_table_caption(document: Document, config: dict[str, Any]) -> bool:
    for paragraph in reversed(document.paragraphs):
        if paragraph.text.strip():
            return paragraph.style.name == "VKR Table Caption" and is_table_caption_text(paragraph.text, config)
    return False


def table_caption_from_header(rows: list[list[str]], table_number: int, config: dict[str, Any]) -> str:
    prefix = config["captions"]["table"]["prefix"]
    separator = config["captions"]["table"]["separator"]
    header = [cell.strip(" `") for cell in rows[0] if cell.strip()] if rows else []
    title = ", ".join(header[:3])
    if len(header) > 3:
        title += " и другие параметры"
    if not title:
        title = "Сводные данные"
    return f"{prefix} {table_number} {separator} {title}"


def add_auto_table_caption(document: Document, rows: list[list[str]], table_number: int, config: dict[str, Any]) -> None:
    if document_ends_with_table_caption(document, config):
        return
    paragraph = document.add_paragraph(style="VKR Table Caption")
    paragraph.text = table_caption_from_header(rows, table_number, config)
    for run in paragraph.runs:
        apply_run_font(run, config)


def consume_table(document: Document, tokens: list[Token], start: int, config: dict[str, Any], table_number: int) -> int:
    rows: list[list[str]] = []
    current_row: list[str] | None = None
    in_cell = False
    i = start + 1
    while i < len(tokens):
        token = tokens[i]
        if token.type == "table_close":
            break
        if token.type == "tr_open":
            current_row = []
        elif token.type == "tr_close" and current_row is not None:
            rows.append(current_row)
            current_row = None
        elif token.type in {"th_open", "td_open"}:
            in_cell = True
        elif token.type in {"th_close", "td_close"}:
            in_cell = False
        elif in_cell and token.type == "inline" and current_row is not None:
            current_row.append(normalize_text(plain_text_from_inline(token).strip(), config))
        i += 1

    if rows:
        add_auto_table_caption(document, rows, table_number, config)
        table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
        table.style = "Table Grid"
        for row_index, row in enumerate(rows):
            for col_index, value in enumerate(row):
                cell = table.cell(row_index, col_index)
                cell.text = value
            if row_index == 0 and config["tables"]["repeat_header_rows"]:
                set_repeat_table_header(table.rows[row_index])
        format_table(table, config, config["document"]["default_text"])
    return i + 1


def add_equation(document: Document, content: str, config: dict[str, Any], equation_number: int) -> None:
    section = document.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin
    number_width = Cm(config["formulas"]["number_column_cm"])
    equation_width = usable_width - number_width

    table = document.add_table(rows=1, cols=2)
    set_table_no_borders(table)
    table.autofit = False
    set_table_preferred_width(table, usable_width)
    table.columns[0].width = int(equation_width)
    table.columns[1].width = int(number_width)

    for cell in table.row_cells(0):
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.left_indent = Cm(0)
            paragraph.paragraph_format.right_indent = Cm(0)

    equation_cell = table.cell(0, 0)
    equation_paragraph = equation_cell.paragraphs[0]
    equation_paragraph.alignment = ALIGNMENTS[config["formulas"]["alignment"]]
    equation_paragraph._p.append(wrap_omath_para(formula_text_to_omml(content)))

    number_cell = table.cell(0, 1)
    number_paragraph = number_cell.paragraphs[0]
    number_paragraph.alignment = ALIGNMENTS[config["formulas"]["number_alignment"]]
    enclosure = config["formulas"]["number_enclosure"]
    number_text = f"({equation_number})" if enclosure == "parentheses" else str(equation_number)
    number_run = number_paragraph.add_run(number_text)
    apply_run_font(number_run, config)


def add_code_block(document: Document, content: str, config: dict[str, Any], equation_number: int | None = None) -> int | None:
    stripped = content.rstrip("\n")
    if looks_like_formula(stripped):
        if equation_number is None:
            raise ValueError("Formula block requires equation numbering")
        add_equation(document, stripped, config, equation_number)
        return equation_number + 1

    paragraph = document.add_paragraph(style="VKR Code Block")
    for line_index, line in enumerate(stripped.splitlines()):
        if line_index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        apply_run_font(run, config)
    return equation_number


def convert_markdown(input_path: Path, output_path: Path, config: dict[str, Any]) -> None:
    if output_path.suffix.lower() != ".docx":
        raise ValueError("Output path must have .docx extension")

    markdown = input_path.read_text(encoding=config["markdown"]["encoding"])
    parser = MarkdownIt("commonmark").enable("table")
    tokens = parser.parse(markdown)

    document = Document()
    configure_document(document, config)

    list_stack: list[str] = []
    ordered_counters: list[int] = []
    table_number = 1
    equation_number = 1
    in_bibliography: list[bool] = [False]
    current_chapter: int | None = None
    chapter_figures: dict[int, int] = {}
    i = 0
    while i < len(tokens):
        token = tokens[i]
        if token.type == "heading_open":
            level = int(token.tag[1])
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                heading_text = plain_text_from_inline(tokens[i + 1]).strip()
                chapter = chapter_from_heading_text(heading_text)
                if chapter is not None and level <= 2:
                    current_chapter = chapter
                if re.match(r"^Список литературы", heading_text, re.IGNORECASE):
                    in_bibliography[0] = True
                add_heading(document, tokens[i + 1], level, config)
                i += 3
                continue
        if token.type == "paragraph_open":
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                inline_token = tokens[i + 1]
                inline_text = plain_text_from_inline(inline_token).strip()
                chapter = chapter_from_heading_text(inline_text)
                if chapter is not None:
                    current_chapter = chapter
                if re.match(r"^Список литературы", inline_text, re.IGNORECASE):
                    in_bibliography[0] = True
                if list_stack and list_stack[-1] == "ordered":
                    ordered_counters[-1] += 1
                    add_paragraph_from_inline(
                        document,
                        inline_token,
                        config,
                        None,
                        ordered_counters[-1],
                        in_bibliography,
                    )
                else:
                    list_style = "List Bullet" if list_stack and list_stack[-1] == "bullet" else None
                    add_paragraph_from_inline(
                        document,
                        inline_token,
                        config,
                        list_style,
                        in_bibliography=in_bibliography,
                    )
                i += 3
                continue
        if token.type == "bullet_list_open":
            list_stack.append("bullet")
        elif token.type == "ordered_list_open":
            list_stack.append("ordered")
            ordered_counters.append(0)
        elif token.type == "bullet_list_close" and list_stack:
            list_stack.pop()
        elif token.type == "ordered_list_close" and list_stack:
            list_stack.pop()
            if ordered_counters:
                ordered_counters.pop()
        elif token.type == "fence":
            fence_info = token.info.strip().lower()
            if fence_info in {"plantuml", "puml"} or (
                fence_info in {"", "text"} and "@startuml" in token.content
            ):
                add_plantuml_figure(document, token.content, config, current_chapter, chapter_figures)
            elif fence_info in {"math", "formula"} or looks_like_formula(token.content.strip()):
                add_equation(document, token.content.strip(), config, equation_number)
                equation_number += 1
            else:
                add_code_block(document, token.content, config)
        elif token.type == "table_open":
            i = consume_table(document, tokens, i, config, table_number)
            table_number += 1
            continue
        elif token.type == "hr":
            document.add_page_break()
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="Input Markdown file")
    parser.add_argument("--config", required=True, type=Path, help="VКР formatting JSON config")
    parser.add_argument("--output", required=True, type=Path, help="Output DOCX file")
    parser.add_argument(
        "--no-assembly",
        action="store_true",
        help="Skip title pages, abstract, and TOC composition",
    )
    args = parser.parse_args()

    config = load_config(args.config)
    assembly = config.get("assembly", {})
    use_assembly = not args.no_assembly and bool(assembly.get("title_pages") or assembly.get("include_toc"))

    if use_assembly:
        with tempfile.TemporaryDirectory() as temp_dir:
            body_path = Path(temp_dir) / "body.docx"
            convert_markdown(args.input, body_path, config)

            abstract_path: Path | None = None
            abstract_md = assembly.get("abstract_md")
            if abstract_md:
                abstract_source = resolve_tool_path(args.config, abstract_md)
                abstract_path = Path(temp_dir) / "abstract.docx"
                convert_markdown(abstract_source, abstract_path, config)

            assemble_vkr_document(body_path, args.output, args.config, config, abstract_path)
    else:
        convert_markdown(args.input, args.output, config)


if __name__ == "__main__":
    main()
